from __future__ import annotations

from pathlib import Path

from docx import Document

from app.schemas.methods import MethodDetail
from app.schemas.results import AnalysisResultPayload


def export_result_to_docx(method: MethodDetail, result: AnalysisResultPayload, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    report_title = f"{method.name}报告"
    document.core_properties.title = report_title
    document.add_heading(report_title, level=0)
    document.add_paragraph(f"任务 ID：{result.task_id}")
    document.add_paragraph(f"方法标识：{result.method_slug}")
    document.add_paragraph(f"结果版本：{result.result_version}")
    document.add_paragraph(f"生成时间：{result.generated_at}")

    document.add_heading("结果摘要", level=1)
    for highlight in result.summary.highlights:
        document.add_paragraph(highlight, style="List Bullet")

    document.add_heading("可读解释", level=1)
    document.add_paragraph("通俗解释")
    document.add_paragraph(result.interpretation.plain_language)
    document.add_paragraph("学术表达")
    document.add_paragraph(result.interpretation.academic_style)

    for table in result.tables:
        document.add_heading(table.title, level=1)
        word_table = document.add_table(rows=1, cols=len(table.columns))
        word_table.style = "Table Grid"
        for index, column in enumerate(table.columns):
            word_table.rows[0].cells[index].text = str(column)

        for row in table.rows:
            cells = word_table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = _format_cell(value)

    if result.charts:
        document.add_heading("图表数据", level=1)
        for chart in result.charts:
            document.add_paragraph(chart.title)
            chart_table = document.add_table(rows=1, cols=2)
            chart_table.style = "Table Grid"
            chart_table.rows[0].cells[0].text = "label"
            chart_table.rows[0].cells[1].text = "value"

            for label, value in zip(chart.data.labels, chart.data.values, strict=True):
                cells = chart_table.add_row().cells
                cells[0].text = label
                cells[1].text = _format_cell(value)

    document.save(destination)


def _format_cell(value: str | int | float | None) -> str:
    if value is None:
        return "-"
    if isinstance(value, float):
        return f"{value:.4f}".rstrip("0").rstrip(".")
    return str(value)
